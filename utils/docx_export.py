from io import BytesIO
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
def make_offer_docx(df, title="Offerte")->bytes:
    doc=Document()
    h=doc.add_heading(title,0); h.alignment=WD_ALIGN_PARAGRAPH.LEFT
    p=doc.add_paragraph("Totaalprijs: "); p.add_run(f"EUR {df['total_cost'].sum():,.2f}").bold=True
    table=doc.add_table(rows=1, cols=7); hdr=table.rows[0].cells
    for i,t in enumerate(["Line","Material","Qty","Mat. cost","Proc. cost","Overhead","Total"]): hdr[i].text=t
    for _,r in df.iterrows():
        row=table.add_row().cells
        row[0].text=str(r.get("line_id","")); row[1].text=str(r.get("material_id",""))
        row[2].text=str(r.get("qty","")); row[3].text=f"{r.get('material_cost',0):,.2f}"
        row[4].text=f"{r.get('process_cost',0):,.2f}"; row[5].text=f"{r.get('overhead',0):,.2f}"
        row[6].text=f"{r.get('total_cost',0):,.2f}"
    bio=BytesIO(); doc.save(bio); return bio.getvalue()
